"""Tests for legal report export tool."""

from __future__ import annotations

from pathlib import Path

import pytest

from oprocess.tools.legal_report_export import (
    _extract_metadata,
    _extract_title,
    _split_sections,
    export_legal_report,
)

DEMO5_PATH = Path("docs/demos/demo-05-jinyuan-legal-dd-report.md")


@pytest.fixture()
def demo5_md() -> str:
    if not DEMO5_PATH.exists():
        pytest.skip("Demo 5 report file not found")
    return DEMO5_PATH.read_text(encoding="utf-8")


@pytest.fixture()
def parsed(demo5_md: str) -> dict:
    return _split_sections(demo5_md)


class TestSplitSections:
    def test_body_has_13_chapters(self, parsed: dict) -> None:
        assert len(parsed["body"]) == 13

    def test_appendix_a_has_items(self, parsed: dict) -> None:
        assert len(parsed["appendix_a"]) >= 13  # 13 MCP + 1 O'Process

    def test_appendix_b_has_13_items(self, parsed: dict) -> None:
        assert len(parsed["appendix_b"]) == 13

    def test_tail_sections_preserved(self, parsed: dict) -> None:
        assert len(parsed["tail_sections"]) == 2

    def test_header_contains_metadata(self, parsed: dict) -> None:
        assert "广州市金元物业" in parsed["header"]
        assert "委托方" in parsed["header"]


class TestNoTechnicalLeakage:
    def test_body_no_mcp(self, parsed: dict) -> None:
        body_text = "\n".join(parsed["body"])
        assert "MCP 流程定位" not in body_text

    def test_body_no_provenance_chain(self, parsed: dict) -> None:
        body_text = "\n".join(parsed["body"])
        assert "provenance_chain" not in body_text

    def test_body_no_search_process(self, parsed: dict) -> None:
        body_text = "\n".join(parsed["body"])
        assert "search_process" not in body_text

    def test_body_no_oprocess_framework(self, parsed: dict) -> None:
        body_text = "\n".join(parsed["body"])
        assert "O'Process 流程框架" not in body_text


class TestAppendixCompleteness:
    def test_appendix_a_contains_mcp(self, parsed: dict) -> None:
        a_text = " ".join(item["content"] for item in parsed["appendix_a"])
        assert "MCP" in a_text

    def test_appendix_b_contains_provenance(self, parsed: dict) -> None:
        b_text = " ".join(item["content"] for item in parsed["appendix_b"])
        assert "provenance_chain" in b_text

    def test_appendix_items_have_chapter_ref(self, parsed: dict) -> None:
        for item in parsed["appendix_a"]:
            assert "chapter" in item
            assert "第" in item["chapter"]


class TestExportEndToEnd:
    def test_generates_docx(self, tmp_path: Path, demo5_md: str) -> None:
        output = tmp_path / "output.docx"
        result = export_legal_report(str(DEMO5_PATH), str(output))
        assert output.exists()
        assert result["chapters"] == 13
        assert result["appendix_a_items"] >= 13
        assert result["appendix_b_items"] == 13

    def test_docx_has_content(self, tmp_path: Path, demo5_md: str) -> None:
        output = tmp_path / "output.docx"
        export_legal_report(str(DEMO5_PATH), str(output))

        from docx import Document

        doc = Document(str(output))
        assert len(doc.paragraphs) > 100
        assert len(doc.tables) > 10

    def test_docx_body_clean(self, tmp_path: Path, demo5_md: str) -> None:
        """Verify the Word body has no technical content."""
        output = tmp_path / "output.docx"
        export_legal_report(str(DEMO5_PATH), str(output))

        from docx import Document

        doc = Document(str(output))
        # Collect text before the appendix section break
        body_text = ""
        for p in doc.paragraphs:
            if "附录A：流程分类节点映射" in p.text:
                break
            body_text += p.text + "\n"

        assert "MCP 流程定位" not in body_text
        assert "provenance_chain" not in body_text


class TestCjkFontStyles:
    def test_normal_style_has_songti(self, tmp_path: Path, demo5_md: str) -> None:
        output = tmp_path / "output.docx"
        export_legal_report(str(DEMO5_PATH), str(output))

        from docx import Document
        from docx.oxml.ns import qn

        doc = Document(str(output))
        normal = doc.styles["Normal"]
        run_props = normal.element.get_or_add_rPr()
        east_asian = run_props.rFonts.get(qn("w:eastAsia"))
        assert east_asian == "宋体"

    def test_heading_styles_have_heiti(self, tmp_path: Path, demo5_md: str) -> None:
        output = tmp_path / "output.docx"
        export_legal_report(str(DEMO5_PATH), str(output))

        from docx import Document
        from docx.oxml.ns import qn

        doc = Document(str(output))
        for level in range(1, 4):
            h = doc.styles[f"Heading {level}"]
            h_run_props = h.element.get_or_add_rPr()
            east_asian = h_run_props.rFonts.get(qn("w:eastAsia"))
            assert east_asian == "黑体", f"Heading {level} should use 黑体"


class TestMetadataExtraction:
    def test_full_extraction(self) -> None:
        header = (
            "# 报告标题\n"
            "> **委托方:** 广东广盐\n"
            "> **目标公司:** 金元物业\n"
            "> **律所:** 北京大成\n"
            "> **基准日:** 2025年3月\n"
            "> **报告出具日:** 2025年4月\n"
        )
        meta = _extract_metadata(header)
        assert meta["委托方"] == "广东广盐"
        assert meta["目标公司"] == "金元物业"
        assert meta["律所"] == "北京大成"
        assert meta["基准日"] == "2025年3月"
        assert meta["报告出具日"] == "2025年4月"

    def test_partial_fields(self) -> None:
        header = "> **委托方:** 测试公司\n普通文本行\n"
        meta = _extract_metadata(header)
        assert meta == {"委托方": "测试公司"}

    def test_empty_header(self) -> None:
        assert _extract_metadata("") == {}

    def test_title_extraction(self) -> None:
        header = "# 法律尽职调查报告\n> **委托方:** X\n"
        assert _extract_title(header) == "法律尽职调查报告"

    def test_title_empty_when_missing(self) -> None:
        assert _extract_title("no heading here") == ""

    def test_split_sections_has_metadata(self, parsed: dict) -> None:
        assert "目标公司" in parsed["metadata"]
        assert "金元物业" in parsed["metadata"]["目标公司"]
        assert parsed["title"] != ""


class TestDachengFormat:
    @pytest.fixture()
    def doc(self, tmp_path: Path, demo5_md: str):
        from docx import Document

        output = tmp_path / "dacheng.docx"
        export_legal_report(str(DEMO5_PATH), str(output))
        return Document(str(output))

    def test_cover_page_exists(self, doc) -> None:
        """Document should have at least 3 sections (cover, body, appendix)."""
        assert len(doc.sections) >= 3

    def test_heading_font_sizes(self, doc) -> None:
        from docx.shared import Pt

        h1 = doc.styles["Heading 1"]
        assert h1.font.size == Pt(16)
        h2 = doc.styles["Heading 2"]
        assert h2.font.size == Pt(14)
        h3 = doc.styles["Heading 3"]
        assert h3.font.size == Pt(12)

    def test_header_footer_present(self, doc) -> None:
        """Body section (index 1) should have header/footer content."""
        body_sec = doc.sections[1]
        paras = body_sec.header.paragraphs
        header_text = paras[0].text if paras else ""
        assert "大成" in header_text

    def test_signature_page(self, doc) -> None:
        """Last paragraphs should contain firm name (signature)."""
        all_text = [p.text for p in doc.paragraphs if p.text.strip()]
        # Signature firm name should appear near the end
        last_20 = all_text[-20:]
        assert any("大成" in t for t in last_20)

    def test_table_header_shading(self, doc) -> None:
        """Content table header row should have shading.

        tables[0] is the cover metadata table (no borders).
        Content tables start from tables[1].
        """
        from docx.oxml.ns import qn

        # Skip cover metadata table, check first content table
        table = doc.tables[1]
        first_cell = table.rows[0].cells[0]
        tc_pr = first_cell._tc.tcPr
        shd = tc_pr.find(qn("w:shd")) if tc_pr is not None else None
        assert shd is not None
        assert shd.get(qn("w:fill")) == "D6E4F0"

    def test_appendix_titles_renamed(self, doc) -> None:
        """Appendix headings should use A/B labels."""
        headings = [
            p.text for p in doc.paragraphs
            if p.style.name.startswith("Heading")
        ]
        appendix_headings = [h for h in headings if "附录" in h]
        assert any("附录A" in h for h in appendix_headings)
        assert any("附录B" in h for h in appendix_headings)

    def test_no_text_marker(self, doc) -> None:
        all_text = [p.text for p in doc.paragraphs]
        assert any("本页以下无正文" in t for t in all_text)
