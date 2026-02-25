"""Dentons cover page builder for legal DD reports."""

from __future__ import annotations

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from oprocess.tools._docx_xml import (
    add_bottom_border,
    remove_table_borders,
    set_run_east_asian,
    setup_page_section,
)


def add_cover_page(
    doc: Document, sections: dict, *, navy_blue, firm_name, firm_url,
) -> None:
    """Create Dentons-style cover page."""
    sec = doc.sections[0]
    setup_page_section(sec)
    sec.header.is_linked_to_previous = False
    sec.footer.is_linked_to_previous = False

    # Firm name (黑体 18pt, navy blue)
    p_firm = doc.add_paragraph()
    p_firm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_firm = p_firm.add_run(firm_name)
    run_firm.font.size = Pt(18)
    run_firm.font.color.rgb = navy_blue
    run_firm.bold = True
    set_run_east_asian(run_firm, "黑体")
    add_bottom_border(p_firm)

    for _ in range(4):
        doc.add_paragraph()

    # Main title (华文中宋 28pt)
    meta = sections.get("metadata", {})
    target = meta.get("目标公司", "")
    title = f"关于{target}之\n法律尽职调查报告" if target else "法律尽职调查报告"
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, line in enumerate(title.split("\n")):
        if i > 0:
            p_title.add_run("\n")
        run_t = p_title.add_run(line)
        run_t.font.size = Pt(28)
        set_run_east_asian(run_t, "华文中宋")

    for _ in range(4):
        doc.add_paragraph()

    _add_cover_metadata(doc, meta)

    for _ in range(2):
        doc.add_paragraph()
    p_url = doc.add_paragraph()
    p_url.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_url = p_url.add_run(firm_url)
    run_url.font.size = Pt(12)
    run_url.font.color.rgb = navy_blue


def _add_cover_metadata(doc: Document, meta: dict[str, str]) -> None:
    """Add metadata table on cover page."""
    fields = [
        ("委托方", meta.get("委托方", "")),
        ("目标公司", meta.get("目标公司", "")),
        ("出具律所", meta.get("律所", "")),
        ("基准日", meta.get("基准日", "")),
        ("报告出具日", meta.get("报告出具日", "")),
    ]
    fields = [(k, v) for k, v in fields if v]
    if not fields:
        return

    table = doc.add_table(rows=len(fields), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for i, (label, value) in enumerate(fields):
        row = table.rows[i]
        _set_cell(row.cells[0], label, Cm(3), WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(row.cells[1], value, Cm(10))
    remove_table_borders(table)


def _set_cell(cell, text: str, width, align=None) -> None:
    """Configure a cover metadata cell."""
    cell.text = text
    cell.width = width
    for p in cell.paragraphs:
        if align:
            p.alignment = align
        for r in p.runs:
            r.font.size = Pt(12)
            set_run_east_asian(r, "宋体")
