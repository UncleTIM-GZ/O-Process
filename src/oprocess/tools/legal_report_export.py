"""Export Demo 5 legal DD report to Word (body/appendix separated).

Parses a Demo 5 markdown report and generates a Word document where:
- Main body contains only business content (事实核查, 法律意见, 风险判断)
- Appendix A contains PCF node mapping (x.1 MCP sections)
- Appendix B contains ProvenanceChain data (x.N 溯源链 sections)
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

# Pre-compiled regex patterns
_RE_CODE_BLOCK = re.compile(r"^```[^\n]*\n.*?\n```", re.MULTILINE | re.DOTALL)
_RE_CODE_PLACEHOLDER = re.compile(r"__CODE_(\d+)__")
_RE_CHAPTER = re.compile(r"^## (.+)$", re.MULTILINE)
_RE_SUBSECTION = re.compile(r"^### (.+)$", re.MULTILINE)
_RE_TABLE_SEP = re.compile(r"^\|[\s:-]+\|$")
_RE_BOLD = re.compile(r"\*\*(.+?)\*\*")
_RE_ORDERED_LIST = re.compile(r"^\d+\.\s")


def export_legal_report(md_path: str, output_path: str) -> dict:
    """Export Demo 5 markdown to Word with body/appendix separation."""
    md = Path(md_path).read_text(encoding="utf-8")
    sections = _split_sections(md)
    doc = _build_docx(sections)
    doc.save(output_path)
    return {
        "chapters": len(sections["body"]),
        "appendix_a_items": len(sections["appendix_a"]),
        "appendix_b_items": len(sections["appendix_b"]),
    }


# ========== Parsing Layer ==========


def _split_sections(md: str) -> dict:
    """Split markdown into body + appendix content."""
    # Stash code blocks to avoid false heading matches
    stash: list[str] = []

    def _stash(m: re.Match) -> str:
        stash.append(m.group(0))
        return f"__CODE_{len(stash) - 1}__"

    cleaned = _RE_CODE_BLOCK.sub(_stash, md)

    # Extract header (everything before first ## chapter)
    first_h2 = _RE_CHAPTER.search(cleaned)
    header = cleaned[: first_h2.start()].strip() if first_h2 else ""

    # Split by H2 headings into [title, content, title, content, ...]
    parts = _RE_CHAPTER.split(cleaned)
    # parts[0] = before first H2, then alternating title/content

    result: dict = {
        "header": _restore(header, stash),
        "body": [],
        "appendix_a": [],
        "appendix_b": [],
        "tail_sections": [],
    }

    for i in range(1, len(parts), 2):
        title = parts[i].strip()
        content = parts[i + 1] if i + 1 < len(parts) else ""
        full_chapter = f"## {title}\n{content}"

        if title.startswith("第") and "章" in title:
            _classify_chapter(title, full_chapter, stash, result)
        elif title.startswith("附录"):
            result["tail_sections"].append(_restore(full_chapter, stash))
        else:
            result["body"].append(_restore(full_chapter, stash))

    return result


def _classify_chapter(
    title: str, chapter_md: str, stash: list[str], result: dict,
) -> None:
    """Classify subsections of a chapter into body vs appendix."""
    # Split by H3 subsections
    sub_parts = _RE_SUBSECTION.split(chapter_md)
    # sub_parts[0] = chapter heading line, then alternating sub_title/sub_content

    body_parts = [sub_parts[0]]  # Always keep chapter heading
    for j in range(1, len(sub_parts), 2):
        sub_title = sub_parts[j].strip()
        sub_content = sub_parts[j + 1] if j + 1 < len(sub_parts) else ""
        full_sub = f"### {sub_title}\n{sub_content}"

        if _is_mcp_section(sub_title):
            result["appendix_a"].append({
                "chapter": title,
                "content": _restore(full_sub, stash),
            })
        elif _is_provenance_section(sub_title):
            result["appendix_b"].append({
                "chapter": title,
                "content": _restore(full_sub, stash),
            })
        elif _is_oprocess_explanation(sub_title):
            result["appendix_a"].append({
                "chapter": title,
                "content": _restore(full_sub, stash),
            })
        else:
            body_parts.append(full_sub)

    result["body"].append(_restore("\n".join(body_parts), stash))


def _is_mcp_section(title: str) -> bool:
    return "MCP" in title and "流程定位" in title


def _is_provenance_section(title: str) -> bool:
    return "溯源链" in title


def _is_oprocess_explanation(title: str) -> bool:
    return "O'Process" in title and "流程框架" in title


def _restore(text: str, stash: list[str]) -> str:
    """Restore stashed code blocks (single-pass regex replacement)."""
    if not stash:
        return text
    return _RE_CODE_PLACEHOLDER.sub(lambda m: stash[int(m.group(1))], text)


# ========== Word Builder Layer ==========


def _build_docx(sections: dict) -> Document:
    """Build Word document from parsed sections."""
    doc = Document()
    _setup_styles(doc)
    _setup_page(doc)

    # Header (title + metadata blockquote)
    _render_markdown(doc, sections["header"])

    # TOC placeholder
    _add_toc(doc)

    # Body chapters
    for chapter_md in sections["body"]:
        _render_markdown(doc, chapter_md)

    # Section break before appendix
    doc.add_section(WD_SECTION.NEW_PAGE)

    # Appendix A: PCF node mapping
    doc.add_heading("附录：流程分类节点映射", level=2)
    for item in sections["appendix_a"]:
        doc.add_heading(item["chapter"], level=3)
        _render_markdown(doc, item["content"])

    # Appendix B: Provenance chains
    doc.add_heading("附录：数据溯源链", level=2)
    for item in sections["appendix_b"]:
        doc.add_heading(item["chapter"], level=3)
        _render_markdown(doc, item["content"])

    # Tail sections (original appendices)
    for tail_md in sections["tail_sections"]:
        _render_markdown(doc, tail_md)

    return doc


def _setup_styles(doc: Document) -> None:
    """Set CJK fonts at style level."""
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run_props = normal.element.get_or_add_rPr()
    run_props.rFonts.set(qn("w:eastAsia"), "宋体")

    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Times New Roman"
        h_run_props = h.element.get_or_add_rPr()
        h_run_props.rFonts.set(qn("w:eastAsia"), "黑体")


def _setup_page(doc: Document) -> None:
    """A4 page with legal document margins."""
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    sec.top_margin, sec.bottom_margin = Cm(2.54), Cm(2.54)
    sec.left_margin, sec.right_margin = Cm(3.17), Cm(3.17)


def _add_toc(doc: Document) -> None:
    """Insert TOC field (user updates via right-click in Word)."""
    p = doc.add_paragraph()
    run = p.add_run()
    r = run._r
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "右键单击此处更新目录"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r.extend([fld_begin, instr, fld_sep, placeholder, fld_end])


def _render_markdown(doc: Document, md: str) -> None:
    """Render markdown text to Word paragraphs."""
    in_table = False
    table_rows: list[list[str]] = []

    for line in md.split("\n"):
        stripped = line.strip()

        if stripped.startswith("## "):
            _flush_table(doc, table_rows, in_table)
            in_table, table_rows = False, []
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            _flush_table(doc, table_rows, in_table)
            in_table, table_rows = False, []
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("|") and stripped.endswith("|"):
            if _RE_TABLE_SEP.match(stripped):
                continue  # Skip separator row
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            table_rows.append(cells)
            in_table = True
        elif in_table and not stripped.startswith("|"):
            _flush_table(doc, table_rows, in_table)
            in_table, table_rows = False, []
            _add_text_line(doc, stripped)
        elif stripped.startswith("```"):
            continue  # Code fence markers handled by stash/restore
        elif stripped.startswith("> "):
            p = doc.add_paragraph(stripped[2:])
            p.paragraph_format.left_indent = Cm(1.0)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif _RE_ORDERED_LIST.match(stripped):
            text = _RE_ORDERED_LIST.sub("", stripped, count=1)
            doc.add_paragraph(text, style="List Number")
        elif stripped == "---":
            continue  # Skip horizontal rules
        elif stripped:
            _add_text_line(doc, stripped)

    _flush_table(doc, table_rows, in_table)


def _flush_table(
    doc: Document, rows: list[list[str]], active: bool,
) -> None:
    """Write accumulated table rows to document."""
    if not active or not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.autofit = False
    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            if ci < ncols:
                table.rows[ri].cells[ci].text = cell_text
    rows.clear()


def _add_text_line(doc: Document, text: str) -> None:
    """Add a text line with bold support."""
    if not text:
        return
    parts = _RE_BOLD.split(text)
    if len(parts) == 1:
        doc.add_paragraph(text)
        return
    p = doc.add_paragraph()
    for k, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        if k % 2 == 1:  # Odd indices are bold captures
            run.bold = True
