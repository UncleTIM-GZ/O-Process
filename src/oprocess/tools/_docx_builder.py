"""Dentons China (大成) format Word builder for legal DD reports.

Generates professional Chinese legal documents with:
- Cover page with firm branding
- Header/footer with firm name and page numbers
- CJK typography (宋体/黑体/华文中宋)
- Table header shading
- Signature page
"""

from __future__ import annotations

import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from oprocess.tools._docx_cover import add_cover_page
from oprocess.tools._docx_xml import (
    add_bottom_border,
    bold_row,
    insert_field,
    set_run_east_asian,
    setup_page_section,
    shade_row,
)

# ========== Constants ==========

NAVY_BLUE = RGBColor(0x00, 0x3D, 0x6B)
LIGHT_BLUE_SHADING = "D6E4F0"
FIRM_NAME = "北京大成律师事务所"
FIRM_URL = "www.dentons.cn"

_RE_TABLE_SEP = re.compile(r"^\|[\s:-]+\|$")
_RE_BOLD = re.compile(r"\*\*(.+?)\*\*")
_RE_ORDERED_LIST = re.compile(r"^\d+\.\s")


def build_docx(sections: dict) -> Document:
    """Build Dentons-format Word document from parsed sections."""
    doc = Document()
    _setup_styles(doc)

    add_cover_page(
        doc, sections,
        navy_blue=NAVY_BLUE, firm_name=FIRM_NAME, firm_url=FIRM_URL,
    )

    body_sec = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_page_section(body_sec)
    _setup_header_footer(body_sec)

    _add_toc(doc)

    for chapter_md in sections["body"]:
        _render_markdown(doc, chapter_md)

    _add_no_text_marker(doc)

    doc.add_section(WD_SECTION.NEW_PAGE)
    _add_appendices(doc, sections)

    for tail_md in sections["tail_sections"]:
        _render_markdown(doc, tail_md)

    _add_signature(doc, sections.get("metadata", {}))
    return doc


# ========== Styles ==========


def _setup_styles(doc: Document) -> None:
    """Configure CJK fonts and heading styles for Dentons format."""
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")

    _cfg = [
        ("Heading 1", 16, Pt(12), Pt(12)),
        ("Heading 2", 14, Pt(6), Pt(6)),
        ("Heading 3", 12, None, None),
    ]
    for name, size, before, after in _cfg:
        h = doc.styles[name]
        h.font.name = "Times New Roman"
        h.font.size = Pt(size)
        h.font.bold = True
        if before:
            h.paragraph_format.space_before = before
        if after:
            h.paragraph_format.space_after = after
        h.element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "黑体")


# ========== Header / Footer ==========


def _setup_header_footer(sec) -> None:
    """Add firm name header with border and page number footer."""
    sec.header.is_linked_to_previous = False
    sec.footer.is_linked_to_previous = False

    hp = (
        sec.header.paragraphs[0]
        if sec.header.paragraphs
        else sec.header.add_paragraph()
    )
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run(FIRM_NAME)
    run.font.size = Pt(9)
    run.font.color.rgb = NAVY_BLUE
    set_run_east_asian(run, "黑体")
    add_bottom_border(hp)

    fp = (
        sec.footer.paragraphs[0]
        if sec.footer.paragraphs
        else sec.footer.add_paragraph()
    )
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run("- ")
    insert_field(fp.add_run()._r, " PAGE ")
    fp.add_run(" -")


# ========== Markdown Rendering ==========


def _render_markdown(doc: Document, md: str) -> None:
    """Render markdown to Word with heading remapping (## → H1)."""
    in_table = False
    table_rows: list[list[str]] = []

    for line in md.split("\n"):
        s = line.strip()

        if s.startswith("#### "):
            in_table = _flush(doc, table_rows, in_table)
            doc.add_heading(s[5:], level=3)
        elif s.startswith("### "):
            in_table = _flush(doc, table_rows, in_table)
            doc.add_heading(s[4:], level=2)
        elif s.startswith("## "):
            in_table = _flush(doc, table_rows, in_table)
            doc.add_heading(s[3:], level=1)
        elif s.startswith("|") and s.endswith("|"):
            if not _RE_TABLE_SEP.match(s):
                cells = [c.strip() for c in s.strip("|").split("|")]
                table_rows.append(cells)
                in_table = True
        elif in_table and not s.startswith("|"):
            in_table = _flush(doc, table_rows, in_table)
            _add_text_line(doc, s)
        elif s.startswith("```") or s == "---":
            continue
        elif s.startswith("> "):
            p = doc.add_paragraph(s[2:])
            p.paragraph_format.left_indent = Cm(1.0)
        elif s.startswith("- "):
            doc.add_paragraph(s[2:], style="List Bullet")
        elif _RE_ORDERED_LIST.match(s):
            doc.add_paragraph(
                _RE_ORDERED_LIST.sub("", s, count=1),
                style="List Number",
            )
        elif s:
            _add_text_line(doc, s)

    _flush(doc, table_rows, in_table)


def _flush(doc: Document, rows: list[list[str]], active: bool) -> bool:
    """Write accumulated table rows with header shading. Returns False."""
    if not active or not rows:
        return False
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.autofit = False
    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            if ci < ncols:
                table.rows[ri].cells[ci].text = cell_text
    if rows:
        shade_row(table.rows[0], LIGHT_BLUE_SHADING)
        bold_row(table.rows[0])
    rows.clear()
    return False


def _add_text_line(doc: Document, text: str) -> None:
    """Add a text line with bold support."""
    if not text:
        return
    parts = _RE_BOLD.split(text)
    if len(parts) == 1:
        doc.add_paragraph(text)
        return
    p = doc.add_paragraph()
    for k, part in enumerate(parts):
        if part:
            run = p.add_run(part)
            if k % 2 == 1:
                run.bold = True


# ========== Appendices & Signature ==========


def _add_appendices(doc: Document, sections: dict) -> None:
    """Add appendix A and B sections."""
    doc.add_heading("附录A：流程分类节点映射", level=1)
    for item in sections["appendix_a"]:
        doc.add_heading(item["chapter"], level=2)
        _render_markdown(doc, item["content"])

    doc.add_heading("附录B：数据溯源链", level=1)
    for item in sections["appendix_b"]:
        doc.add_heading(item["chapter"], level=2)
        _render_markdown(doc, item["content"])


def _add_no_text_marker(doc: Document) -> None:
    """Add centered '（本页以下无正文）' marker."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("（本页以下无正文）")


def _add_toc(doc: Document) -> None:
    """Insert TOC field (user updates via right-click in Word)."""
    r = doc.add_paragraph().add_run()._r
    insert_field(r, 'TOC \\o "1-3" \\h \\z \\u', "右键单击此处更新目录")


def _add_signature(doc: Document, meta: dict[str, str]) -> None:
    """Add signature page with firm name and date."""
    doc.add_section(WD_SECTION.NEW_PAGE)
    for _ in range(8):
        doc.add_paragraph()

    firm = meta.get("律所", FIRM_NAME)
    p_firm = doc.add_paragraph()
    p_firm.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_firm = p_firm.add_run(firm)
    run_firm.font.size = Pt(12)
    set_run_east_asian(run_firm, "宋体")

    date_str = meta.get("报告出具日", "")
    if date_str:
        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_date = p_date.add_run(date_str)
        run_date.font.size = Pt(12)
        set_run_east_asian(run_date, "宋体")
